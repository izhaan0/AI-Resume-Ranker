"""
Resume parsing service.

Extracts raw text from PDF and DOCX files using:
  - pdfplumber (primary PDF parser)
  - PyMuPDF / fitz (PDF fallback)
  - python-docx (DOCX)
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── Raw text extraction ───────────────────────────────────────────────────────

def extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF file using pdfplumber, falling back to PyMuPDF."""
    text = ""

    # Primary: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages.append(page_text)
            text = "\n".join(pages)
        if text.strip():
            return text
    except Exception as exc:
        logger.warning("pdfplumber failed for %s: %s. Trying PyMuPDF.", path.name, exc)

    # Fallback: PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages = [page.get_text("text") for page in doc]
        doc.close()
        text = "\n".join(pages)
    except Exception as exc:
        logger.error("PyMuPDF also failed for %s: %s", path.name, exc)

    return text


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", path.name, exc)
        return ""


def extract_text(path: Path) -> str:
    """Route to the correct extractor based on file extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── Name extraction ───────────────────────────────────────────────────────────

def extract_name(text: str) -> Optional[str]:
    """
    Heuristic: the candidate's name is usually one of the first non-empty lines
    and looks like 'Firstname Lastname' (2-4 words, no digits).
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:8]:
        # Skip lines that look like headers or contact info
        if any(kw in line.lower() for kw in ["resume", "curriculum", "cv", "@", "http", "linkedin", "phone"]):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() and w.isalpha() for w in words):
            return line
    return None


# ── Email extraction ──────────────────────────────────────────────────────────

def extract_email(text: str) -> Optional[str]:
    pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else None


# ── Phone extraction ──────────────────────────────────────────────────────────

def extract_phone(text: str) -> Optional[str]:
    pattern = r"(\+?\d[\d\s\-().]{7,}\d)"
    match = re.search(pattern, text)
    if match:
        phone = re.sub(r"[^\d+]", "", match.group(0))
        if 7 <= len(phone) <= 15:
            return match.group(0).strip()
    return None


# ── Experience extraction ─────────────────────────────────────────────────────

_EXPERIENCE_PATTERNS = [
    r"(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience",
    r"experience\s*[:\-]?\s*(\d+)\+?\s*(?:years?|yrs?)",
    r"(\d+)\+?\s*(?:years?|yrs?)\s+(?:in|of|with)",
]

def extract_experience_years(text: str) -> Optional[float]:
    lower = text.lower()
    for pattern in _EXPERIENCE_PATTERNS:
        m = re.search(pattern, lower)
        if m:
            return float(m.group(1))

    # Fallback: count work experience date ranges like "2018 – 2022"
    date_ranges = re.findall(r"(\d{4})\s*[–\-—to]+\s*(\d{4}|present|current)", lower)
    if date_ranges:
        import datetime
        current_year = datetime.date.today().year
        total = 0.0
        for start, end in date_ranges:
            s = int(start)
            e = current_year if end in ("present", "current") else int(end)
            if 1980 <= s <= current_year and s <= e:
                total += e - s
        if total > 0:
            return round(min(total, 50), 1)

    return None


# ── Education extraction ──────────────────────────────────────────────────────

_DEGREES = [
    "phd", "ph.d", "doctorate", "m.tech", "m.sc", "m.s.", "msc", "ms",
    "mba", "m.b.a", "b.tech", "b.sc", "b.s.", "bsc", "bs", "bachelor",
    "master", "associate", "diploma",
]

def extract_education(text: str) -> Optional[str]:
    lower = text.lower()
    for degree in _DEGREES:
        if degree in lower:
            # Return the sentence/line that contains the degree
            for line in text.splitlines():
                if degree in line.lower() and len(line) < 200:
                    return line.strip()
    return None


# ── Job titles extraction ─────────────────────────────────────────────────────

_TITLE_KEYWORDS = [
    "engineer", "developer", "scientist", "analyst", "architect", "manager",
    "lead", "director", "consultant", "specialist", "designer", "researcher",
    "intern", "associate", "senior", "junior", "staff", "principal",
]

def extract_job_titles(text: str) -> list[str]:
    titles = []
    for line in text.splitlines():
        lower = line.lower()
        if any(kw in lower for kw in _TITLE_KEYWORDS) and len(line.split()) <= 8:
            cleaned = line.strip()
            if cleaned and cleaned not in titles:
                titles.append(cleaned)
        if len(titles) >= 5:
            break
    return titles
